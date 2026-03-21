"""Build a final DOCX transcript from multiple rendered UFM templates."""

from __future__ import annotations

import re
import shutil
import textwrap
import uuid
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from .context_builder import ContextBuilder
from .template_registry import get_template_path
from .template_renderer import TemplateRenderer
from .template_selector import TemplateSelector
from .ufm_formatter import UFMFormatter


class DocumentBuilder:
    """Render selected templates and merge them into one DOCX output."""

    BODY_LINES_PER_PAGE = 25
    BODY_WIDTHS = {
        "Q": 54,
        "A": 54,
        "COLLOQUY": 48,
        "PAREN": 48,
        "SECTION": 60,
        "PLAIN": 60,
    }
    DEFAULT_FINALIZATION_OPTIONS = {
        "show_format_box": True,
        "show_line_numbers": True,
        "show_header": True,
        "show_footer": True,
        "header_text": "CAUSE NO. 2024-XXXX",
        "footer_text": "Page 1",
    }

    def build_document(
        self,
        job_data: dict,
        output_path: str,
        finalization_options: dict | None = None,
    ) -> None:
        """Build a final transcript document from template-driven sections.

        Args:
            job_data: Raw job information used to build template context and
                select applicable templates.
            output_path: Final DOCX path for the assembled transcript.
            finalization_options: Optional toggle-based finalization controls.

        Raises:
            FileNotFoundError: If a required template file is missing.
            RuntimeError: If rendering, merging, or saving the document fails.
            ValueError: If upstream context or selection validation fails.
        """
        output_file = Path(output_path)
        temp_dir = Path("temp_render") / uuid.uuid4().hex
        rendered_files: list[Path] = []

        try:
            context = ContextBuilder().build_context(job_data)
            template_names = TemplateSelector().select_templates(job_data)
            if not template_names:
                raise ValueError("No templates selected for document build.")
            print(f"Templates selected: {template_names}")

            temp_dir.mkdir(parents=True, exist_ok=True)
            renderer = TemplateRenderer()

            for template_name in template_names:
                template_path = Path(get_template_path(template_name))
                is_optional = template_name in {"signature_block", "certification_cont"}
                if not template_path.is_file():
                    if is_optional:
                        print(f"Skipping optional missing template: {template_name}")
                        continue
                    raise FileNotFoundError(f"Template file not found: {template_path}")

                temp_file = temp_dir / f"{template_name}_{uuid.uuid4().hex}.docx"
                print(f"Rendering {template_name}...")
                renderer.render_template(
                    template_path=str(template_path),
                    context=context,
                    output_path=str(temp_file),
                )
                print(f"Rendered: {temp_file}")
                rendered_files.append(temp_file)

            builder_options = self._normalize_finalization_options(finalization_options)
            transcript_data = self._get_transcript_data(job_data)
            self.build_transcript_body(
                transcript_data,
                show_line_numbers=bool(builder_options["show_line_numbers"]),
                show_format_box=bool(builder_options["show_format_box"]),
            )
            transcript_body = self._transcript_body_document

            print("Merging sections...")
            ordered_sections = self._build_ordered_sections(template_names, rendered_files, transcript_body)
            if not ordered_sections:
                raise ValueError("No document sections were generated for merge.")

            merged_document = ordered_sections[0]
            for section_document in ordered_sections[1:]:
                merged_document.add_page_break()
                self._append_document(merged_document, section_document)

            output_file.parent.mkdir(parents=True, exist_ok=True)
            formatter = UFMFormatter()
            formatter.enforce_document(merged_document)
            if builder_options["show_header"]:
                self.apply_header(merged_document, job_data)
            if builder_options["show_footer"]:
                self.apply_footer(merged_document)
            merged_document.save(str(output_file))
            print(f"Output saved: {output_file}")
        except FileNotFoundError:
            raise
        except ValueError:
            raise
        except Exception as exc:
            raise RuntimeError(f"Document build failed: {exc}") from exc
        finally:
            if temp_dir.exists():
                shutil.rmtree(temp_dir, ignore_errors=True)

    def build_transcript_body(
        self,
        transcript_data: list[dict],
        show_line_numbers: bool = True,
        show_format_box: bool = True,
    ) -> list[Paragraph]:
        """Build transcript body paragraphs in code and retain the source document."""
        self._transcript_body_document = self._build_transcript_body_document(
            transcript_data,
            show_line_numbers=show_line_numbers,
            show_format_box=show_format_box,
        )
        return list(self._transcript_body_document.paragraphs)

    def paginate(self, transcript_data: list[dict]) -> list[list[dict]]:
        """Split transcript data into 25-line pages based on wrapped line estimates."""
        pages: list[list[dict]] = []
        current_page: list[dict] = []
        line_count = 0

        for entry in transcript_data:
            wrapped_lines = self._wrap_entry(entry)
            required_lines = max(1, len(wrapped_lines))

            if current_page and line_count + required_lines > self.BODY_LINES_PER_PAGE:
                pages.append(current_page)
                current_page = []
                line_count = 0

            current_page.append({**entry, "_wrapped_lines": wrapped_lines})
            line_count += required_lines

        if current_page:
            pages.append(current_page)

        print("Pagination complete")
        return pages

    @staticmethod
    def _append_document(target: Document, source: Document) -> None:
        """Append the body elements from one document into another."""
        for element in list(source.element.body):
            # Section properties are handled by the target document.
            if element.tag.endswith("sectPr"):
                continue
            target.element.body.append(deepcopy(element))

    def _build_ordered_sections(
        self,
        template_names: list[str],
        rendered_files: list[Path],
        transcript_body: Document,
    ) -> list[Document]:
        """Create the final document section order with code-generated transcript body."""
        rendered_map = {
            template_name: rendered_file
            for template_name, rendered_file in zip(template_names, rendered_files, strict=False)
        }
        ordered_sections: list[Document] = []
        body_inserted = False

        for template_name in template_names:
            rendered_file = rendered_map.get(template_name)
            if rendered_file is None:
                continue

            ordered_sections.append(Document(str(rendered_file)))

            if template_name in {"witness_setup_remote", "witness_setup_standard"} and not body_inserted:
                ordered_sections.append(transcript_body)
                body_inserted = True

        if not body_inserted:
            ordered_sections.append(transcript_body)

        return ordered_sections

    def _build_transcript_body_document(
        self,
        transcript_data: list[dict],
        show_line_numbers: bool,
        show_format_box: bool,
    ) -> Document:
        """Generate the transcript body entirely in code."""
        print("Generating transcript body...")
        body_document = Document()
        formatter = UFMFormatter()
        formatter.apply_document_format(body_document)

        for element in list(body_document.element.body):
            body_document.element.body.remove(element)

        pages = self.paginate(transcript_data)
        for page_index, page in enumerate(pages):
            used_lines = 0

            for entry in page:
                wrapped_lines = entry.get("_wrapped_lines") or self._wrap_entry(entry)
                entry_type = str(entry.get("type", "PLAIN")).upper()
                speaker_label = str(entry.get("speaker_label", "")).strip()

                for line_index, line in enumerate(wrapped_lines):
                    paragraph = body_document.add_paragraph()
                    formatter.apply_paragraph_format(paragraph)
                    line_number = used_lines + 1
                    self._apply_body_indent(paragraph, entry_type)
                    self._apply_body_tab_stops(paragraph, entry_type)
                    self._apply_body_line_spacing(paragraph)
                    self._apply_body_borders(
                        paragraph,
                        line_number=line_number,
                        show_format_box=show_format_box,
                    )

                    run_text = self._format_body_line(
                        entry_type,
                        line,
                        speaker_label,
                        line_index,
                        len(wrapped_lines),
                        line_number,
                        show_line_numbers,
                    )
                    run = paragraph.add_run(run_text)
                    formatter.apply_font(run)
                    used_lines += 1

            while used_lines < self.BODY_LINES_PER_PAGE:
                paragraph = body_document.add_paragraph()
                formatter.apply_paragraph_format(paragraph)
                line_number = used_lines + 1
                self._apply_body_tab_stops(paragraph, "PLAIN")
                self._apply_body_line_spacing(paragraph)
                self._apply_body_borders(
                    paragraph,
                    line_number=line_number,
                    show_format_box=show_format_box,
                )
                run = paragraph.add_run(f"{line_number:>2} " if show_line_numbers else "")
                formatter.apply_font(run)
                used_lines += 1

            if used_lines != self.BODY_LINES_PER_PAGE:
                raise RuntimeError("UFM violation: page does not contain exactly 25 lines")

            if page_index < len(pages) - 1:
                paragraph = body_document.add_paragraph()
                paragraph.add_run().add_break(WD_BREAK.PAGE)

        return body_document

    def _get_transcript_data(self, job_data: dict) -> list[dict]:
        """Return transcript entries from job data or build a deterministic fallback body."""
        transcript_data = job_data.get("transcript_data")
        if isinstance(transcript_data, list) and transcript_data:
            return transcript_data

        witness_name = str(job_data.get("witness_name", "the witness")).strip() or "the witness"
        return [
            {"type": "Q", "text": "Please state your full name for the record."},
            {"type": "A", "text": witness_name},
            {
                "type": "Q",
                "text": "Did you review the exhibits and deposition notice before appearing today?",
            },
            {
                "type": "A",
                "text": "Yes. I reviewed the notice, the scheduling email, and the documents that were attached before the deposition started.",
            },
            {
                "type": "COLLOQUY",
                "speaker_label": "MR. SMITH",
                "text": "Let the record reflect that the witness is appearing remotely.",
            },
            {"type": "PAREN", "text": "Discussion off the record"},
            {
                "type": "Q",
                "text": "Are all of your answers today based on your personal knowledge unless you say otherwise?",
            },
            {
                "type": "A",
                "text": "Yes -- and if I do not know or remember something, I will say so.",
            },
        ]

    def _wrap_entry(self, entry: dict) -> list[str]:
        """Wrap one transcript entry to deterministic UFM-compatible line widths."""
        entry_type = str(entry.get("type", "PLAIN")).upper()
        text = self._post_process_text(str(entry.get("text", "")).strip())
        width = self.BODY_WIDTHS.get(entry_type, self.BODY_WIDTHS["PLAIN"])

        if not text:
            return [""]

        return textwrap.wrap(
            text,
            width=width,
            break_long_words=False,
            break_on_hyphens=False,
        ) or [text]

    @staticmethod
    def _post_process_text(text: str) -> str:
        """Normalize transcript punctuation for deterministic body rendering."""
        processed = text.replace("—", " -- ").replace("–", " -- ")
        processed = re.sub(r"([.!?])\s+", r"\1  ", processed)
        return processed

    @staticmethod
    def _apply_body_indent(paragraph: Paragraph, entry_type: str) -> None:
        """Apply type-specific paragraph indent behavior for transcript body lines."""
        fmt = paragraph.paragraph_format
        if entry_type in {"COLLOQUY", "PAREN"}:
            fmt.left_indent = Inches(1.5)
            fmt.first_line_indent = Inches(0)
        else:
            fmt.left_indent = Inches(0)
            fmt.first_line_indent = Inches(0)

    @staticmethod
    def _apply_body_tab_stops(paragraph: Paragraph, entry_type: str) -> None:
        """Apply explicit tab stops required by UFM body formatting."""
        tab_stops = paragraph.paragraph_format.tab_stops
        tab_stops.clear_all()

        if entry_type in {"Q", "A"}:
            tab_stops.add_tab_stop(Inches(0.3), WD_TAB_ALIGNMENT.LEFT)
            tab_stops.add_tab_stop(Inches(0.5), WD_TAB_ALIGNMENT.LEFT)
            tab_stops.add_tab_stop(Inches(1.0), WD_TAB_ALIGNMENT.LEFT)
            return

        if entry_type in {"COLLOQUY", "PAREN"}:
            tab_stops.add_tab_stop(Inches(0.3), WD_TAB_ALIGNMENT.LEFT)
            tab_stops.add_tab_stop(Inches(1.5), WD_TAB_ALIGNMENT.LEFT)
            return

        tab_stops.add_tab_stop(Inches(0.3), WD_TAB_ALIGNMENT.LEFT)
        tab_stops.add_tab_stop(Inches(0.5), WD_TAB_ALIGNMENT.LEFT)

    @staticmethod
    def _apply_body_line_spacing(paragraph: Paragraph) -> None:
        """Force exact line spacing on transcript-body paragraphs."""
        fmt = paragraph.paragraph_format
        fmt.line_spacing_rule = 4  # WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(28)
        paragraph.style.font.name = "Courier New"
        paragraph.style.font.size = Pt(12)

    @staticmethod
    def _apply_body_borders(paragraph: Paragraph, line_number: int, show_format_box: bool) -> None:
        """Apply format-box borders to transcript body paragraphs when enabled."""
        if not show_format_box:
            return

        p_pr = paragraph._p.get_or_add_pPr()
        for child in list(p_pr):
            if child.tag == qn("w:pBdr"):
                p_pr.remove(child)

        p_bdr = OxmlElement("w:pBdr")

        border_names = ["left", "right"]
        if line_number == 1:
            border_names.append("top")
        if line_number == DocumentBuilder.BODY_LINES_PER_PAGE:
            border_names.append("bottom")

        for border_name in border_names:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "6")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
            p_bdr.append(border)

        p_pr.append(p_bdr)

    @staticmethod
    def _normalize_finalization_options(finalization_options: dict | None) -> dict:
        """Accept both builder-style and finalizer-style option keys."""
        options = dict(DocumentBuilder.DEFAULT_FINALIZATION_OPTIONS)
        if finalization_options:
            options.update(finalization_options)

        if "apply_box" in options:
            options["show_format_box"] = bool(options["apply_box"])
        if "apply_line_numbers" in options:
            options["show_line_numbers"] = bool(options["apply_line_numbers"])
        if "apply_header_footer" in options:
            apply_header_footer = bool(options["apply_header_footer"])
            options["show_header"] = apply_header_footer
            options["show_footer"] = apply_header_footer

        return options

    def apply_header(self, document: Document, job_data: dict) -> None:
        """Apply a centered header on page 2+ with cause number, case style, and page field."""
        cause_number = str(job_data.get("cause_number") or "2024-XXXX").strip()
        case_style = str(job_data.get("case_style") or job_data.get("case_name") or "").strip()

        for section in document.sections:
            section.different_first_page_header_footer = True
            header = section.header
            if not header.paragraphs:
                header.add_paragraph()

            for paragraph in header.paragraphs:
                paragraph.clear()

            header_para = header.paragraphs[0]
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_para.add_run(f"CAUSE NO. {cause_number}\n")
            if case_style:
                header_para.add_run(f"{case_style}\n")
            header_para.add_run("Page ")
            self._append_page_number_field(header_para)

    def apply_footer(self, document: Document) -> None:
        """Apply a centered footer with continuous page numbering."""
        for section in document.sections:
            footer = section.footer
            if not footer.paragraphs:
                footer.add_paragraph()

            for paragraph in footer.paragraphs:
                paragraph.clear()

            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.add_run("Page ")
            self._append_page_number_field(footer_para)

    @staticmethod
    def _append_page_number_field(paragraph: Paragraph) -> None:
        """Append a live PAGE field to the provided paragraph."""
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "

        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")

        text = OxmlElement("w:t")
        text.text = "1"

        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")

        run = paragraph.add_run()
        run._r.append(begin)
        run._r.append(instr_text)
        run._r.append(separate)
        run._r.append(text)
        run._r.append(end)

    @staticmethod
    def _format_body_line(
        entry_type: str,
        line: str,
        speaker_label: str,
        line_index: int,
        total_lines: int,
        line_number: int,
        show_line_numbers: bool,
    ) -> str:
        """Render one transcript body line using UFM tab conventions."""
        line_prefix = f"{line_number:>2} " if show_line_numbers else ""
        if entry_type == "Q":
            return f"{line_prefix}\tQ.\t{line}" if line_index == 0 else f"{line_prefix}\t{line}"
        if entry_type == "A":
            return f"{line_prefix}\tA.\t{line}" if line_index == 0 else f"{line_prefix}\t{line}"
        if entry_type == "COLLOQUY":
            return f"{line_prefix}\t{speaker_label.upper()}:  {line}" if line_index == 0 else f"{line_prefix}\t{line}"
        if entry_type == "PAREN":
            if total_lines == 1:
                return f"{line_prefix}\t({line})"
            if line_index == 0:
                return f"{line_prefix}\t({line}"
            if line_index == total_lines - 1:
                return f"{line_prefix}\t{line})"
            return f"{line_prefix}\t{line}"
        if entry_type == "SECTION":
            return f"{line_prefix}{line.upper()}"
        return f"{line_prefix}{line}"
