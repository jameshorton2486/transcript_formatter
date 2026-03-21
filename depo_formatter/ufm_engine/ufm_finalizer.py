"""Final compliance-layer adjustments for completed UFM transcript documents."""

from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class UFMFinalizer:
    """Apply optional court-specific finishing rules to a completed document."""

    def apply_headers_footers(
        self,
        doc,
        header_text: str,
        footer_text: str,
        show_header: bool,
        show_footer: bool,
    ) -> None:
        """Apply header/footer text with first-page header suppression."""
        for section in doc.sections:
            section.different_first_page_header_footer = True
            header = section.header
            footer = section.footer

            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

            header_para.clear()
            footer_para.clear()
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if show_header and header_text:
                header_para.add_run(header_text)

            if show_footer and footer_text:
                footer_para.add_run(footer_text)

    def apply_page_box(self, doc) -> None:
        """Apply a page border box to each section via direct XML manipulation."""
        for section in doc.sections:
            sect_pr = section._sectPr
            self._remove_existing_child(sect_pr, "w:pgBorders")

            pg_borders = OxmlElement("w:pgBorders")
            pg_borders.set(qn("w:offsetFrom"), "page")

            for border_name in ("top", "left", "bottom", "right"):
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "12")
                border.set(qn("w:space"), "24")
                border.set(qn("w:color"), "000000")
                pg_borders.append(border)

            sect_pr.append(pg_borders)

    def apply_line_numbering(self, doc) -> None:
        """Apply line numbering metadata to each section via direct XML manipulation."""
        for section in doc.sections:
            sect_pr = section._sectPr
            self._remove_existing_child(sect_pr, "w:lnNumType")

            line_numbering = OxmlElement("w:lnNumType")
            line_numbering.set(qn("w:countBy"), "1")
            line_numbering.set(qn("w:start"), "1")
            line_numbering.set(qn("w:distance"), "240")

            sect_pr.append(line_numbering)

    def finalize_document(self, doc, options: dict) -> None:
        """Apply optional finalization features based on provided toggle settings."""
        if options.get("show_header") or options.get("show_footer") or options.get("apply_header_footer"):
            self.apply_headers_footers(
                doc,
                options.get("header_text", "").strip(),
                options.get("footer_text", "").strip(),
                show_header=bool(options.get("show_header", options.get("apply_header_footer", False))),
                show_footer=bool(options.get("show_footer", options.get("apply_header_footer", False))),
            )

        if options.get("apply_box"):
            self.apply_page_box(doc)

        if options.get("apply_line_numbers"):
            self.apply_line_numbering(doc)

    @staticmethod
    def _remove_existing_child(sect_pr, tag_name: str) -> None:
        """Remove an existing section child element before re-applying it."""
        existing = sect_pr.find(qn(tag_name))
        if existing is not None:
            sect_pr.remove(existing)
