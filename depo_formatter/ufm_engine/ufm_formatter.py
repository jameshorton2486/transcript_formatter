"""Strict post-merge formatting utilities for UFM transcript documents."""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt


class UFMFormatter:
    """Apply strict UFM formatting rules to a DOCX document."""

    def apply_document_format(self, doc: Document) -> None:
        """Apply page geometry to each document section."""
        for section in doc.sections:
            section.left_margin = Inches(1.75)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)

    def apply_paragraph_format(self, paragraph) -> None:
        """Apply spacing, indent, and tab stop rules to a paragraph."""
        fmt = paragraph.paragraph_format
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(28)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.left_indent = Inches(0)
        fmt.first_line_indent = Inches(0)

        fmt.tab_stops.clear_all()
        fmt.tab_stops.add_tab_stop(Inches(0.5), WD_TAB_ALIGNMENT.LEFT)
        fmt.tab_stops.add_tab_stop(Inches(1.0), WD_TAB_ALIGNMENT.LEFT)
        fmt.tab_stops.add_tab_stop(Inches(1.5), WD_TAB_ALIGNMENT.LEFT)

    def apply_font(self, run) -> None:
        """Apply the required transcript font to a run."""
        run.font.name = "Courier New"
        run.font.size = Pt(12)

    def format_qa(self, doc: Document, lines: list[str]) -> None:
        """Add Q/A-formatted paragraphs to a document."""
        for line in lines:
            p = doc.add_paragraph()
            self.apply_paragraph_format(p)

            if line.startswith("Q."):
                run = p.add_run("Q.\t" + line[2:].strip())
            elif line.startswith("A."):
                run = p.add_run("A.\t" + line[2:].strip())
            else:
                run = p.add_run(line.strip())

            self.apply_font(run)

    def format_speaker(self, doc: Document, speaker: str, text: str) -> None:
        """Add a formatted speaker line to a document."""
        p = doc.add_paragraph()
        self.apply_paragraph_format(p)
        run = p.add_run("\t\t\t" + speaker.upper() + ":  " + text)
        self.apply_font(run)

    def format_parenthetical(self, doc: Document, text: str) -> None:
        """Add a formatted parenthetical line to a document."""
        p = doc.add_paragraph()
        self.apply_paragraph_format(p)
        run = p.add_run("\t\t\t(" + text + ")")
        self.apply_font(run)

    def apply_interruptions(self, text: str) -> str:
        """Normalize interruption punctuation to double dashes."""
        return text.replace(" - ", " -- ")

    def enforce_document(self, doc: Document) -> None:
        """Apply page, paragraph, font, and interruption rules to a document."""
        self.apply_document_format(doc)

        for paragraph in doc.paragraphs:
            self.apply_paragraph_format(paragraph)

            for run in paragraph.runs:
                run.text = self.apply_interruptions(run.text)
                self.apply_font(run)
