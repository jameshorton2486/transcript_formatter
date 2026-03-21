from __future__ import annotations

import json
import re
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path

import pdfplumber
from docx import Document


SPEAKER_LABEL_PATTERN = re.compile(r"^Speaker\s+(\d):?\s*$", re.IGNORECASE)
ABBREVIATIONS = re.compile(
    r"\b(Dr|Mr|Mrs|Ms|Jr|Sr|vs|No|Vol|Dept|Corp|Inc|Ltd|P\.C|PLLC|St|Ave|Blvd)\.$",
    re.IGNORECASE,
)
ATTORNEY_KEYWORDS = re.compile(
    r"i\s+am\s+\w+.*i\s+represent|my\s+name\s+is\s+\w+.*counsel|"
    r"representing\s+the\s+(plaintiff|defendant)|counsel\s+for\s+the\s+(plaintiff|defendant)",
    re.IGNORECASE,
)
REPORTER_KEYWORDS = re.compile(
    r"raise\s+your\s+right\s+hand|do\s+you\s+solemnly|"
    r"swear\s+or\s+affirm|spell\s+your\s+name|the\s+reporter",
    re.IGNORECASE,
)
VIDEOGRAPHER_KEYWORDS = re.compile(
    r"we\s+are\s+(now\s+)?on\s+(the\s+)?record|"
    r"going\s+off\s+(the\s+)?record|this\s+is\s+the\s+video",
    re.IGNORECASE,
)


@dataclass
class SpeakerMap:
    witness: int | None = None
    lead_attorney: int | None = None
    opposing_attorney: int | None = None
    court_reporter: int | None = None
    videographer: int | None = None
    unassigned: list[int] = field(default_factory=list)


@dataclass
class LoadedTranscript:
    text: str
    source_type: str
    blocks: list[tuple[int, list[str]]] = field(default_factory=list)
    speaker_map: SpeakerMap | None = None


def load_transcript(path: str) -> LoadedTranscript:
    file_path = Path(path)
    suffix = file_path.suffix.lower()

    if suffix == ".txt":
        return LoadedTranscript(text=load_txt(path), source_type="txt")
    if suffix == ".json":
        return load_json(path)
    if suffix == ".docx":
        return load_docx_transcript(path)
    if suffix == ".pdf":
        return LoadedTranscript(text=load_pdf(path), source_type="pdf")

    raise ValueError("Unsupported file type. Please select a .txt, .json, .docx, or .pdf file.")


def load_file(path: str) -> str:
    return load_transcript(path).text


def load_txt(path: str) -> str:
    return Path(path).read_text(encoding="utf-8").strip()


def load_docx(path: str) -> str:
    return load_docx_transcript(path).text


def load_json(path: str) -> LoadedTranscript:
    payload = json.loads(Path(path).read_text(encoding="utf-8"))

    blocks = _extract_json_speaker_blocks(payload)
    if blocks:
        return LoadedTranscript(
            text=render_blocks(blocks),
            source_type="json",
            blocks=blocks,
            speaker_map=build_speaker_map(blocks, scan_limit=30),
        )

    paragraph_text = _extract_json_paragraph_text(payload)
    if paragraph_text:
        return LoadedTranscript(
            text=paragraph_text,
            source_type="json",
            blocks=[],
            speaker_map=None,
        )

    raw_transcript = _extract_json_raw_transcript(payload)
    if raw_transcript:
        return LoadedTranscript(
            text=raw_transcript,
            source_type="json",
            blocks=[],
            speaker_map=None,
        )

    raise ValueError("The JSON file did not contain usable transcript text.")


def load_docx_transcript(path: str) -> LoadedTranscript:
    blocks, speaker_map = parse_deepgram_docx(path)
    if blocks:
        return LoadedTranscript(
            text=render_blocks(blocks),
            source_type="docx",
            blocks=blocks,
            speaker_map=speaker_map,
        )

    document = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return LoadedTranscript(
        text="\n".join(paragraphs).strip(),
        source_type="docx",
        blocks=[],
        speaker_map=speaker_map,
    )


def load_pdf(path: str) -> str:
    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text.strip())

    return "\n\n".join(pages).strip()


def _extract_json_speaker_blocks(payload: object) -> list[tuple[int, list[str]]]:
    utterances = _find_first_list(payload, ("utterances",))
    if isinstance(utterances, list):
        blocks: list[tuple[int, list[str]]] = []
        for item in utterances:
            if not isinstance(item, dict):
                continue
            speaker = item.get("speaker")
            transcript = item.get("transcript") or item.get("text")
            if speaker is None or not isinstance(transcript, str) or not transcript.strip():
                continue
            try:
                speaker_id = int(speaker)
            except (TypeError, ValueError):
                continue
            blocks.append(assemble_block(speaker_id, [transcript]))
        if blocks:
            return blocks

    paragraphs = _find_first_list(payload, ("paragraphs",))
    if isinstance(paragraphs, list):
        fragments_by_speaker: dict[int, list[str]] = {}
        ordered_speakers: list[int] = []
        for item in paragraphs:
            if not isinstance(item, dict):
                continue
            speaker = item.get("speaker")
            text = _paragraph_text_from_json(item)
            if speaker is None or not text:
                continue
            try:
                speaker_id = int(speaker)
            except (TypeError, ValueError):
                continue
            if speaker_id not in fragments_by_speaker:
                fragments_by_speaker[speaker_id] = []
                ordered_speakers.append(speaker_id)
            fragments_by_speaker[speaker_id].append(text)
        if fragments_by_speaker:
            return [
                assemble_block(speaker_id, fragments_by_speaker[speaker_id])
                for speaker_id in ordered_speakers
            ]

    return []


def _extract_json_paragraph_text(payload: object) -> str:
    paragraphs = _find_first_list(payload, ("paragraphs",))
    if not isinstance(paragraphs, list):
        return ""

    extracted: list[str] = []
    for item in paragraphs:
        if not isinstance(item, dict):
            continue
        text = _paragraph_text_from_json(item)
        if text:
            extracted.append(text)
    return "\n\n".join(extracted).strip()


def _extract_json_raw_transcript(payload: object) -> str:
    transcript = _find_first_string(payload, ("transcript",))
    return transcript.strip() if transcript else ""


def _paragraph_text_from_json(item: dict) -> str:
    if isinstance(item.get("text"), str) and item["text"].strip():
        return item["text"].strip()
    sentences = item.get("sentences")
    if isinstance(sentences, list):
        parts = [
            sentence.get("text", "").strip()
            for sentence in sentences
            if isinstance(sentence, dict) and sentence.get("text", "").strip()
        ]
        if parts:
            return " ".join(parts).strip()
    return ""


def _find_first_list(payload: object, keys: tuple[str, ...]) -> list | None:
    if isinstance(payload, dict):
        for key, value in payload.items():
            if key in keys and isinstance(value, list):
                return value
            nested = _find_first_list(value, keys)
            if nested is not None:
                return nested
    elif isinstance(payload, list):
        for item in payload:
            nested = _find_first_list(item, keys)
            if nested is not None:
                return nested
    return None


def _find_first_string(payload: object, keys: tuple[str, ...]) -> str | None:
    if isinstance(payload, dict):
        for key, value in payload.items():
            if key in keys and isinstance(value, str) and value.strip():
                return value
            nested = _find_first_string(value, keys)
            if nested is not None:
                return nested
    elif isinstance(payload, list):
        for item in payload:
            nested = _find_first_string(item, keys)
            if nested is not None:
                return nested
    return None


def detect_speaker_label(paragraph_text: str) -> int | None:
    match = SPEAKER_LABEL_PATTERN.match(paragraph_text.strip())
    if match:
        return int(match.group(1))
    return None


def split_into_paragraphs(text: str) -> list[str]:
    raw_parts = re.split(r"(?<=[.?!])\s+(?=[A-Z])", text.strip())

    paragraphs: list[str] = []
    for part in raw_parts:
        part = part.strip()
        if not part:
            continue
        if paragraphs and ABBREVIATIONS.search(paragraphs[-1]):
            paragraphs[-1] = f"{paragraphs[-1]} {part}"
        else:
            paragraphs.append(part)

    return paragraphs


def assemble_block(speaker_id: int, fragments: list[str]) -> tuple[int, list[str]]:
    joined = " ".join(fragments)
    joined = re.sub(r" +", " ", joined).strip()

    paragraphs = split_into_paragraphs(joined)
    if not paragraphs:
        paragraphs = [joined] if joined else []

    return speaker_id, paragraphs


def build_speaker_map(blocks: list[tuple[int, list[str]]], scan_limit: int = 30) -> SpeakerMap:
    sample = blocks[:scan_limit]
    block_counts = Counter(speaker_id for speaker_id, _ in sample)

    speaker_text: dict[int, str] = {}
    for speaker_id, paragraphs in sample:
        speaker_text.setdefault(speaker_id, "")
        speaker_text[speaker_id] += " " + " ".join(paragraphs)

    speaker_map = SpeakerMap()
    assigned: set[int] = set()

    for sid, text in speaker_text.items():
        if REPORTER_KEYWORDS.search(text) and speaker_map.court_reporter is None:
            speaker_map.court_reporter = sid
            assigned.add(sid)
        elif VIDEOGRAPHER_KEYWORDS.search(text) and speaker_map.videographer is None:
            speaker_map.videographer = sid
            assigned.add(sid)
        elif ATTORNEY_KEYWORDS.search(text):
            if speaker_map.lead_attorney is None:
                speaker_map.lead_attorney = sid
                assigned.add(sid)
            elif speaker_map.opposing_attorney is None:
                speaker_map.opposing_attorney = sid
                assigned.add(sid)

    remaining = [(sid, count) for sid, count in block_counts.most_common() if sid not in assigned]
    for sid, _count in remaining:
        if speaker_map.witness is None:
            speaker_map.witness = sid
            assigned.add(sid)
        elif speaker_map.lead_attorney is None:
            speaker_map.lead_attorney = sid
            assigned.add(sid)
        elif speaker_map.opposing_attorney is None:
            speaker_map.opposing_attorney = sid
            assigned.add(sid)
        elif speaker_map.videographer is None:
            speaker_map.videographer = sid
            assigned.add(sid)
        else:
            speaker_map.unassigned.append(sid)

    return speaker_map


def parse_deepgram_docx(docx_path: str) -> tuple[list[tuple[int, list[str]]], SpeakerMap]:
    doc = Document(docx_path)

    blocks: list[tuple[int, list[str]]] = []
    current_speaker: int | None = None
    current_fragments: list[str] = []

    for para in doc.paragraphs:
        text = para.text
        speaker_id = detect_speaker_label(text)

        if speaker_id is not None:
            if current_speaker is not None and current_fragments:
                blocks.append(assemble_block(current_speaker, current_fragments))
            current_speaker = speaker_id
            current_fragments = []
        elif current_speaker is not None:
            current_fragments.append(text)

    if current_speaker is not None and current_fragments:
        blocks.append(assemble_block(current_speaker, current_fragments))

    speaker_map = build_speaker_map(blocks, scan_limit=30)
    return blocks, speaker_map


def render_blocks(blocks: list[tuple[int, list[str]]]) -> str:
    rendered_blocks: list[str] = []
    for speaker_id, paragraphs in blocks:
        block_lines = [f"Speaker {speaker_id}:"]
        block_lines.extend(paragraphs)
        rendered_blocks.append("\n".join(line for line in block_lines if line.strip()))
    return "\n\n".join(rendered_blocks).strip()
